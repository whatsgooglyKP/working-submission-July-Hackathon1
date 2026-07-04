import os
import re
import logging
import traceback
from fastapi import APIRouter, HTTPException, status, Response
from api.schemas import ApplicationRequest, APIOrchestratorResult, ApplicationStrategy, ResumeTailorRequest, ResumeTailorResponse
from api.pdf_generator import generate_strategy_pdf
from job_scraper import search_linkedin_jobs, get_linkedin_job_description
from agents.orchestrator import EasyApplierOrchestrator

logger = logging.getLogger("easyapplier.api")

router = APIRouter()

# Resolve Environment Variables
GCP_PROJECT_ID = os.getenv("GCP_PROJECT_ID", "ringed-land-397222")
APPHUB_APPLICATION = os.getenv("APPHUB_APPLICATION", "easyapplier")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")

def extract_linkedin_profile_details(url: str) -> dict:
    if not url:
        return {
            "name": "Kevin Scott Pinard",
            "headline": "AI Developer, Data Analyst, & Automation Architect",
            "location": "Orlando, FL",
            "skills": ["Python", "AI", "SQL", "BigQuery", "Data Analysis", "Automation"],
            "raw_slug": ""
        }
    
    # Extract path portion
    path_part = ""
    if "/in/" in url:
        path_part = url.split("/in/")[-1]
    elif "/pub/" in url:
        path_part = url.split("/pub/")[-1]
    else:
        path_part = url.split("/")[-1] or url.split("/")[-2] if len(url.split("/")) > 1 else url
        
    path_part = path_part.split("?")[0].split("#")[0].strip()
    
    # Normalize path portion
    normalized = re.sub(r'[_/]', '-', path_part)
    parts = [p.strip() for p in normalized.split("-") if p.strip()]
    
    # Filter out numeric or hexadecimal IDs at the end
    cleaned_parts = []
    for p in parts:
        if re.match(r'^[0-9a-fA-F]+$', p) and (len(p) >= 6 or p.isdigit()):
            continue
        cleaned_parts.append(p)
        
    if not cleaned_parts:
        cleaned_parts = ["pinard", "kevin"] # Fallback if empty slug
        
    # Lookups for classification
    known_locations = {
        "orlando": "Orlando, FL",
        "nyc": "New York, NY",
        "newyork": "New York, NY",
        "sf": "San Francisco, CA",
        "sanfrancisco": "San Francisco, CA",
        "london": "London, UK",
        "austin": "Austin, TX",
        "seattle": "Seattle, WA",
        "boston": "Boston, MA",
        "chicago": "Chicago, IL",
        "miami": "Miami, FL",
        "atlanta": "Atlanta, GA",
        "florida": "Florida, USA",
        "california": "California, USA",
        "texas": "Texas, USA"
    }
    
    known_roles = {
        "developer": "Developer",
        "engineer": "Engineer",
        "scientist": "Scientist",
        "architect": "Architect",
        "analyst": "Analyst",
        "consultant": "Consultant",
        "manager": "Manager",
        "lead": "Lead",
        "specialist": "Specialist",
        "expert": "Expert"
    }
    
    known_skills = {
        "python": "Python",
        "ai": "AI",
        "ml": "Machine Learning",
        "data": "Data Science/Analytics",
        "cloud": "Cloud Computing",
        "aws": "AWS",
        "sql": "SQL",
        "react": "React",
        "java": "Java",
        "automation": "Automation",
        "agentic": "Agentic AI",
        "agent": "AI Agents",
        "bi": "Business Intelligence",
        "devops": "DevOps",
        "tableau": "Tableau",
        "bigquery": "BigQuery",
        "scrum": "Agile/Scrum",
        "pm": "Project Management",
        "fullstack": "Full Stack",
        "frontend": "Frontend",
        "backend": "Backend"
    }
    
    name_parts = []
    location_list = []
    headline_parts = []
    skills_list = []
    
    for i, word in enumerate(cleaned_parts):
        w_lower = word.lower()
        if w_lower in known_locations:
            location_list.append(known_locations[w_lower])
        elif w_lower in known_roles:
            headline_parts.append(known_roles[w_lower])
            skills_list.append(known_roles[w_lower])
        elif w_lower in known_skills:
            skills_list.append(known_skills[w_lower])
            headline_parts.append(known_skills[w_lower])
        else:
            if len(name_parts) < 3 and not any(char.isdigit() for char in word):
                name_parts.append(word.title())
            else:
                skills_list.append(word.title())
                
    if len(name_parts) == 1:
        name_str = name_parts[0]
        if name_str.lower() == "pinardkevin":
            name_parts = ["Kevin", "Scott", "Pinard"]
        elif name_str.lower() == "kevinpinard":
            name_parts = ["Kevin", "Pinard"]
        else:
            name_parts = [name_str.title()]
            
    name = " ".join(name_parts) if name_parts else "Kevin Scott Pinard"
    location = location_list[0] if location_list else "Orlando, FL"
    
    if headline_parts:
        seen = set()
        clean_headline_parts = []
        for hp in headline_parts:
            if hp not in seen:
                seen.add(hp)
                clean_headline_parts.append(hp)
        headline = " | ".join(clean_headline_parts)
    else:
        headline = "AI Developer, Data Analyst, & Automation Architect"
        
    if not skills_list:
        skills_list = ["Python", "AI", "SQL", "BigQuery", "Data Analysis", "Automation"]
        
    skills_list = list(dict.fromkeys(skills_list))
    
    return {
        "name": name,
        "headline": headline,
        "location": location,
        "skills": skills_list,
        "raw_slug": path_part
    }

def get_default_candidate_resume() -> str:
    """Helper to load a default candidate resume from local build documents or fallback info."""
    profile_parts = []
    path1 = "C:/Users/pinar/Kaggle-AI-Agents-Course/Google Kaggle 5DAYAI Vibe Coding Project/RESUME.Google.Day3.docx"
    path2 = "C:/Users/pinar/Google Final Submission Materials/Resume Base Build.docx"
    
    if os.path.exists(path1):
        try:
            import docx
            doc = docx.Document(path1)
            t = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if t:
                profile_parts.append(f"--- CANDIDATE DAY 3 PROFILE ---\n{t}")
        except Exception:
            pass
            
    if os.path.exists(path2):
        try:
            import docx
            doc = docx.Document(path2)
            t = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if t:
                profile_parts.append(f"--- CANDIDATE BASE BUILD PROFILE ---\n{t}")
        except Exception:
            pass
            
    if profile_parts:
        return "\n\n".join(profile_parts)
    
    # Static realistic fallback profile for the system
    return """
    Kevin Scott Pinard
    Email: Kevinpolymath@gmail.com
    Phone: 352-406-3847
    Location: Orlando, FL
    Portfolio: https://www.linkedin.com/in/pinardkevin | https://github.com/kevinpolymath
    
    Summary:
    Results-driven AI Developer, Data Analyst, and Automation Architect. Expert at building agentic workflows, parsing high-dimensional dataset collections, and deploying enterprise-grade cloud integrations.
    
    Skills:
    - Programming: Python, SQL, JavaScript, HTML/CSS, PowerShell, Bash
    - AI & ML: LLMs (Gemini, GPT), LangChain, Vertex AI, Pydantic, Agents SDK, RAG
    - Data & Cloud: BigQuery, PostgreSQL, Tableau, Google Cloud Platform (Cloud Run, GCS), Docker, Git
    
    Experience:
    - Founder / AI Automation Consultant at Aligned Labs (2024 - Present):
      * Built agentic automation architectures to optimize business operations, achieving 40% reduction in manual data tasks.
      * Developed customized LLM applications using Google Vertex AI and Pydantic validation frameworks.
    - Data Architect / Technical Analyst (2021 - 2024):
      * Developed complex BigQuery SQL pipelines for dashboard analytics and predictive forecasting.
      * Designed end-to-end data ingestion flows handling millions of records daily.
    """

@router.get("/health")
async def health_check():
    """Liveness probe utilized by cloud scaling services like Google Cloud Run."""
    return {
        "status": "healthy",
        "service": "easyapplier-agent-system",
        "model_configured": GEMINI_MODEL
    }

@router.get("/api/info")
async def info():
    """Metadata regarding GCP registration and workspace deployment."""
    return {
        "gcp_project_id": GCP_PROJECT_ID,
        "apphub_application": APPHUB_APPLICATION,
        "region": "us-east1"
    }

@router.get("/api/jobs")
async def get_linkedin_jobs(title: str, limit: int = 10):
    """Fetches recently posted job listings on LinkedIn guest search."""
    if not title:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Job title query parameter 'title' is required."
        )
    try:
        jobs = search_linkedin_jobs(title, limit=limit)
        return jobs
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error scraping LinkedIn jobs: {str(e)}"
        )

@router.get("/api/jobs/description")
async def get_job_description(url: str):
    """Retrieves full scraped LinkedIn job description for a given LinkedIn URL."""
    if not url:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Query parameter 'url' is required."
        )
    try:
        desc = get_linkedin_job_description(url)
        if not desc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Could not scrape job description. Make sure the URL is a valid guest LinkedIn job link."
            )
        return {"description": desc}
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error fetching job description: {str(e)}"
        )

@router.post("/api/apply", response_model=ApplicationStrategy)
async def analyze_application(request: ApplicationRequest):
    """
    Launches the full multi-agent coordination workflow:
    1. Job Recon (ResearcherAgent) -> JobSpecification
    2. Candidate Profiler (ResearcherAgent) -> CandidateProfile
    3. Strategic Gap Analysis (StrategistAgent) -> GapAnalysis
    4. Copywriting Tailor (TailorAgent) -> TailorOutput
    
    This fully evaluates candidate fit and generates highly optimized
    application strategies, tailored resumes, customized cover letters, 
    and target interview prep questions.
    """
    logger.info("Starting Multi-Agent Orchestrator workflow for job: %s", request.job_title)

    # Use supplied resume or fall back to default profile
    resume_text_to_use = request.resume_text
    
    # Check if a LinkedIn URL was passed instead of resume text
    is_url = False
    if resume_text_to_use:
        cleaned_url = resume_text_to_use.strip().lower()
        if (cleaned_url.startswith("http://") or cleaned_url.startswith("https://")) and "linkedin.com" in cleaned_url:
            is_url = True
            
    if not resume_text_to_use or is_url:
        logger.info("Synthesizing Candidate Profile from LinkedIn URL and default profile")
        url_to_use = resume_text_to_use if is_url else "https://www.linkedin.com/in/pinardkevin"
        profile_details = extract_linkedin_profile_details(url_to_use)
        default_resume = get_default_candidate_resume()
        
        resume_text_to_use = f"""
        =====================================================
        EXTRACTED CANDIDATE PROFILE (LINKEDIN URL)
        =====================================================
        Full Name: {profile_details['name']}
        Professional Headline: {profile_details['headline']}
        Target Location: {profile_details['location']}
        Inferred Skills & Focus: {", ".join(profile_details['skills'])}
        
        =====================================================
        CANDIDATE BASE RESUME (FALLBACK & EXPERIENCE SOURCE)
        =====================================================
        {default_resume}
        
        CRITICAL PARSING GUIDELINES:
        - The candidate's name MUST be parsed as "{profile_details['name']}".
        - The candidate's summary and location MUST align with "{profile_details['location']}" and headline "{profile_details['headline']}".
        - Merge the inferred skills ({", ".join(profile_details['skills'])}) into the resume.
        - Ensure all tailored outputs use "{profile_details['name']}" as the candidate's name.
        """

    # If job description contains a URL, scrape full details dynamically
    job_description_to_use = request.job_description
    link_match = re.search(r'https?://[^\s]+', job_description_to_use)
    if link_match:
        try:
            logger.info("Scraping full LinkedIn job description from URL")
            scraped_desc = get_linkedin_job_description(link_match.group(0))
            if scraped_desc:
                job_description_to_use = scraped_desc
        except Exception as e:
            logger.warning("Failed to scrape LinkedIn job description: %s", str(e))

    try:
        # Run the official Multi-Agent Orchestrator
        logger.info("Running parallelized EasyApplier Multi-Agent suite...")
        orchestrator = EasyApplierOrchestrator()
        result = orchestrator.run_workflow(
            job_title=request.job_title,
            raw_job_description=job_description_to_use,
            raw_resume_text=resume_text_to_use,
            user_notes=request.user_notes
        )
        
        # Map OrchestratorResult into ApplicationStrategy
        mapped_strategy = ApplicationStrategy(
            match_score=result.gap_analysis.match_score,
            fit_summary=result.gap_analysis.fit_summary,
            cover_letter=result.tailor_output.cover_letter,
            tailored_resume=result.tailor_output.tailored_resume,
            resume_suggestions=result.gap_analysis.gap_areas,
            interview_prep=result.tailor_output.interview_prep
        )
        
        return mapped_strategy

    except Exception as e:
        tb = traceback.format_exc()
        logger.error("Multi-Agent workflow failed with exception traceback:\n%s", tb)
        
        err_msg = str(e).upper()
        if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg or "RATE_LIMIT" in err_msg:
            raise HTTPException(
                status_code=status.HTTP_429_TOO_MANY_REQUESTS,
                detail="Rate limit reached. Please wait a moment and try again."
            )
            
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Multi-Agent execution failed: {str(e)}"
        )

from pydantic import BaseModel, Field

class PDFGenerationPayload(BaseModel):
    strategy: ApplicationStrategy
    job_title: str
    company: str
    linkedin_url: str

@router.post("/api/apply/pdf")
async def get_compiled_pdf(payload: PDFGenerationPayload):
    """
    Compiles the tailored resume, cover letter, alignment score, suggestions,
    and interview prep into a premium, beautifully formatted multi-page PDF career dossier.
    """
    logger.info("Compiling tailored Career Path dossier PDF for: %s", payload.job_title)
    try:
        pdf_bytes = generate_strategy_pdf(
            res=payload.strategy,
            job_title=payload.job_title,
            company=payload.company,
            linkedin_url=payload.linkedin_url
        )
        
        filename = f"EasyApplier_Tailored_Career_Package_{payload.company.replace(' ', '_')}.pdf"
        
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
    except Exception as e:
        logger.error("Failed compiling PDF package: %s", str(e))
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"PDF generation failed: {str(e)}"
        )

@router.post("/api/apply/resume", response_model=ResumeTailorResponse)
async def tailor_resume_endpoint(request: ResumeTailorRequest):
    """
    Tailor a resume specifically based on job posting basic/preferred qualifications
    and matching it to the candidate's LinkedIn URL.
    Returns ONLY the tailored resume Markdown.
    """
    logger.info("Tailoring resume specifically for LinkedIn profile: %s", request.linkedin_url)
    
    # Extract candidate profile details
    profile_details = extract_linkedin_profile_details(request.linkedin_url)
    default_resume = get_default_candidate_resume()
    
    resume_text_to_use = f"""
    =====================================================
    EXTRACTED CANDIDATE PROFILE (LINKEDIN URL)
    =====================================================
    Full Name: {profile_details['name']}
    Professional Headline: {profile_details['headline']}
    Target Location: {profile_details['location']}
    Inferred Skills & Focus: {", ".join(profile_details['skills'])}
    
    =====================================================
    CANDIDATE BASE RESUME (FALLBACK & EXPERIENCE SOURCE)
    =====================================================
    {default_resume}
    
    CRITICAL PARSING GUIDELINES:
    - The candidate's name MUST be parsed as "{profile_details['name']}".
    - The candidate's summary and location MUST align with "{profile_details['location']}" and headline "{profile_details['headline']}".
    - Merge the inferred skills ({", ".join(profile_details['skills'])}) into the resume.
    - Ensure all tailored outputs use "{profile_details['name']}" as the candidate's name.
    """

    job_description_to_use = request.job_url_or_text
    link_match = re.search(r'https?://[^\s]+', job_description_to_use)
    if link_match:
        try:
            logger.info("Scraping full LinkedIn job description from URL")
            scraped_desc = get_linkedin_job_description(link_match.group(0))
            if scraped_desc:
                job_description_to_use = scraped_desc
        except Exception as e:
            logger.warning("Failed to scrape LinkedIn job description: %s", str(e))

    job_title = request.job_title or "Target Position"

    try:
        orchestrator = EasyApplierOrchestrator()
        result = orchestrator.run_workflow(
            job_title=job_title,
            raw_job_description=job_description_to_use,
            raw_resume_text=resume_text_to_use,
            user_notes="Focus strictly on basic and preferred qualifications to align the resume."
        )
        return ResumeTailorResponse(tailored_resume=result.tailor_output.tailored_resume)
    except Exception as e:
        tb = traceback.format_exc()
        logger.error("Resume tailoring failed with exception traceback:\n%s", tb)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Resume tailoring failed: {str(e)}"
        )

