import sys
import io
import os

# Failsafe: Ensure standard streams use UTF-8 encoding on Windows without stream replacement
if sys.stdout is not None:
    try:
        sys.stdout.reconfigure(encoding='utf-8')
    except Exception:
        pass
if sys.stderr is not None:
    try:
        sys.stderr.reconfigure(encoding='utf-8')
    except Exception:
        pass

import streamlit as st
import pandas as pd
import plotly.express as px
import requests
from google import genai
from google.genai import types
from google.genai.errors import APIError
from google.oauth2.credentials import Credentials
from job_scraper import search_linkedin_jobs, get_linkedin_job_description
from pydantic import BaseModel, Field
from dotenv import load_dotenv
import docx
import pypdf
from io import BytesIO
from agents.orchestrator import EasyApplierOrchestrator
from fpdf import FPDF

def clean_for_pdf(text: str) -> str:
    if not text:
        return ""
    # Map common non-latin-1/fancy unicode characters to safe equivalents
    replacements = {
        '\u2018': "'", '\u2019': "'",  # Smart single quotes
        '\u201c': '"', '\u201d': '"',  # Smart double quotes
        '\u2013': '-', '\u2014': '-',  # En/em dashes
        '\u2022': '-', '\u25c6': '-',  # Bullets
        '\u200b': '',                  # Zero-width space
        '\u2112': 'L', '\u2115': 'N',
    }
    for orig, rep in replacements.items():
        text = text.replace(orig, rep)
    
    # Encode to latin-1 and ignore anything that can't be encoded
    try:
        return text.encode('latin-1', 'replace').decode('latin-1')
    except Exception:
        return text.encode('ascii', 'ignore').decode('ascii')

class CareerPathPDF(FPDF):
    def header(self):
        if self.page_no() > 1:
            self.set_font('Helvetica', 'B', 8)
            self.set_text_color(26, 86, 219) # Accent Blue
            self.cell(0, 5, 'EASYAPPLIER CAREER DOSSIER', border=0, ln=0, align='L')
            
            self.set_font('Helvetica', '', 8)
            self.set_text_color(107, 114, 128) # Slate grey
            self.cell(0, 5, 'Tailored Application Strategy', border=0, ln=1, align='R')
            
            # Subtle header line
            self.set_draw_color(229, 231, 235)
            self.set_linewidth(0.4)
            self.line(20, self.get_y(), 190, self.get_y())
            self.ln(5)

    def footer(self):
        if self.page_no() > 1:
            self.set_y(-15)
            self.set_font('Helvetica', 'I', 8)
            self.set_text_color(156, 163, 175)
            self.set_draw_color(229, 231, 235)
            self.line(20, self.get_y() - 2, 190, self.get_y() - 2)
            
            self.cell(0, 10, 'CONFIDENTIAL ◆ AI-OPTIMIZED ATS COMPLIANT DOSSIER', border=0, ln=0, align='L')
            self.cell(0, 10, f'Page {self.page_no()}', border=0, ln=1, align='R')

def extract_candidate_name(resume_text: str) -> str:
    if not resume_text:
        return "Kevin Scott Pinard"
    lines = [line.strip() for line in resume_text.split("\n") if line.strip()]
    if lines:
        for line in lines:
            line_clean = line.replace("#", "").replace("*", "").strip()
            # Find first line that looks like a name (usually less than 40 chars and doesn't start with a section heading)
            if len(line_clean) < 40 and not any(h in line_clean.lower() for h in ["summary", "skills", "experience", "education", "profile", "focus"]):
                return line_clean
    return "Kevin Scott Pinard"

def check_page_break(pdf, space_needed: float = 20):
    if pdf.get_y() + space_needed > pdf.h - pdf.b_margin:
        pdf.add_page()

def write_rich_text(pdf, text: str, default_font="Helvetica", default_size=9.5, default_style="", default_color=(31, 41, 55), line_height=5.5):
    pdf.set_font(default_font, default_style, default_size)
    pdf.set_text_color(*default_color)
    
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:
            pdf.set_font(default_font, default_style + "B" if "B" not in default_style else default_style, default_size)
        else:
            pdf.set_font(default_font, default_style.replace("B", ""), default_size)
        pdf.write(line_height, clean_for_pdf(part))

def write_markdown_line_by_line(pdf, text: str):
    if not text:
        return
    
    lines = text.split("\n")
    for line in lines:
        line_stripped = line.strip()
        
        # Horizontal rule
        if line_stripped in ("---", "***", "___"):
            pdf.ln(3)
            pdf.set_draw_color(209, 213, 219)
            pdf.set_linewidth(0.4)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(3)
            continue
            
        # Heading 1: `# Text`
        if line_stripped.startswith("# "):
            check_page_break(pdf, 25)
            pdf.ln(5)
            cleaned_line = line_stripped[2:].strip()
            write_rich_text(pdf, cleaned_line, default_font="Helvetica", default_size=15, default_style="B", default_color=(26, 86, 219), line_height=7)
            pdf.ln(8)
            continue
            
        # Heading 2: `## Text`
        elif line_stripped.startswith("## "):
            check_page_break(pdf, 20)
            pdf.ln(4)
            cleaned_line = line_stripped[3:].strip()
            write_rich_text(pdf, cleaned_line, default_font="Helvetica", default_size=12, default_style="B", default_color=(17, 24, 39), line_height=6)
            pdf.ln(6)
            
            # Section underline
            pdf.set_draw_color(229, 231, 235)
            pdf.set_linewidth(0.4)
            pdf.line(pdf.l_margin, pdf.get_y() - 1, pdf.w - pdf.r_margin, pdf.get_y() - 1)
            pdf.ln(2.5)
            continue
            
        # Heading 3: `### Text`
        elif line_stripped.startswith("### "):
            check_page_break(pdf, 15)
            pdf.ln(3)
            cleaned_line = line_stripped[4:].strip()
            write_rich_text(pdf, cleaned_line, default_font="Helvetica", default_size=10, default_style="B", default_color=(55, 65, 81), line_height=5.5)
            pdf.ln(5.5)
            continue
            
        # Bullet list: `- Text` or `* Text`
        elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
            pdf.ln(0.5)
            pdf.set_font("Helvetica", "", 9.5)
            pdf.set_text_color(31, 41, 55)
            pdf.cell(5, 5.5, "-", ln=0)
            
            cleaned_line = line_stripped[2:].strip()
            
            old_l_margin = pdf.l_margin
            pdf.set_left_margin(old_l_margin + 5)
            write_rich_text(pdf, cleaned_line, default_font="Helvetica", default_size=9.5, default_style="", default_color=(31, 41, 55), line_height=5.5)
            pdf.ln(5.5)
            pdf.set_left_margin(old_l_margin)
            continue
            
        # Blockquote: `> Text`
        elif line_stripped.startswith("> "):
            pdf.ln(1.5)
            cleaned_line = line_stripped[2:].strip()
            old_l_margin = pdf.l_margin
            
            x_pos = pdf.get_x()
            y_start = pdf.get_y()
            
            pdf.set_left_margin(old_l_margin + 6)
            write_rich_text(pdf, cleaned_line, default_font="Helvetica", default_size=9.5, default_style="I", default_color=(75, 85, 99), line_height=5.5)
            pdf.ln(5.5)
            
            y_end = pdf.get_y()
            pdf.set_draw_color(26, 86, 219)
            pdf.set_linewidth(0.8)
            pdf.line(x_pos + 1, y_start, x_pos + 1, y_end - 1.5)
            
            pdf.set_left_margin(old_l_margin)
            continue
            
        # Standard line or empty line
        else:
            if not line_stripped:
                pdf.ln(2.5)
                continue
                
            write_rich_text(pdf, line_stripped, default_font="Helvetica", default_size=9.5, default_style="", default_color=(31, 41, 55), line_height=5.5)
            pdf.ln(5.5)

def generate_strategy_pdf(res, job_title: str, company: str, linkedin_url: str) -> bytes:
    # Set up PDF with premium geometry
    pdf = CareerPathPDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(20, 20, 20) # 20mm margins are much more elegant than 15mm
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.alias_nb_pages()
    pdf.add_page()
    
    # Extract candidate name dynamically
    candidate_name = extract_candidate_name(res.tailored_resume)
    
    # ----------------------------------------------------
    # COVER PAGE (Page 1)
    # ----------------------------------------------------
    pdf.ln(10)
    
    # Decorative header element: Elegant horizontal brand stripe
    pdf.set_fill_color(26, 86, 219) # Brand Blue
    pdf.rect(20, 25, 170, 4, style="F")
    
    pdf.ln(20)
    
    # Large beautiful typography
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(26, 86, 219) # Brand Blue
    pdf.cell(0, 6, "AI-OPTIMIZED CAREER PREPARATION DOSSIER", ln=1, align="L")
    pdf.ln(2)
    
    pdf.set_font("Helvetica", "B", 28)
    pdf.set_text_color(17, 24, 39) # Elegant near-black
    pdf.cell(0, 11, "Tailored Career Path", ln=1, align="L")
    pdf.cell(0, 11, "Document Package", ln=1, align="L")
    pdf.ln(8)
    
    # Decorative subtitle line
    pdf.set_draw_color(229, 231, 235)
    pdf.set_linewidth(1)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(10)
    
    # Beautiful presentation subtitle
    pdf.set_font("Helvetica", "", 10.5)
    pdf.set_text_color(75, 85, 99)
    pdf.multi_cell(0, 6, "A highly optimized professional portfolio containing a custom tailored resume, personalized cover letter, strategic keywords/gap analysis, and targeted interview preparation instructions.")
    pdf.ln(15)
    
    # Metadata block - layout like a clean table
    pdf.set_fill_color(249, 250, 251) # soft grey card background
    pdf.set_draw_color(229, 231, 235)
    pdf.set_linewidth(0.4)
    pdf.rect(20, pdf.get_y(), 170, 45, style="DF")
    
    y_meta = pdf.get_y() + 5
    
    def draw_meta_row(label, val, y_offset, is_link=False):
        pdf.set_xy(25, y_offset)
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(107, 114, 128) # muted grey
        pdf.cell(45, 6, label, ln=0)
        
        pdf.set_font("Helvetica", "B" if not is_link else "", 9.5)
        pdf.set_text_color(26, 86, 219) if is_link else pdf.set_text_color(31, 41, 55)
        if is_link:
            pdf.cell(0, 6, clean_for_pdf(val), ln=1, link=val)
        else:
            pdf.cell(0, 6, clean_for_pdf(val), ln=1)
            
    draw_meta_row("PREPARED FOR:", candidate_name, y_meta)
    draw_meta_row("TARGET POSITION:", job_title, y_meta + 8)
    draw_meta_row("TARGET COMPANY:", company, y_meta + 16)
    draw_meta_row("PORTFOLIO SOURCE:", linkedin_url, y_meta + 24, is_link=True)
    from datetime import datetime
    draw_meta_row("GENERATED ON:", datetime.now().strftime("%B %d, %Y"), y_meta + 32)
    
    # Move cursor past the card
    pdf.set_xy(20, y_meta + 45)
    pdf.ln(12)
    
    # Beautiful ATS Match Card
    pdf.set_draw_color(229, 231, 235)
    pdf.set_fill_color(249, 250, 251)
    pdf.rect(20, pdf.get_y(), 170, 26, style="DF")
    
    # Draw left indicator border reflecting match level
    if res.match_score >= 80:
        accent_rgb = (16, 185, 129) # Success Green
    elif res.match_score >= 60:
        accent_rgb = (245, 158, 11) # Amber
    else:
        accent_rgb = (239, 68, 68) # Red
        
    pdf.set_fill_color(*accent_rgb)
    pdf.rect(20, pdf.get_y(), 3, 26, style="F")
    
    # Render match score text
    pdf.set_xy(27, pdf.get_y() + 5)
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(107, 114, 128)
    pdf.cell(100, 5, "ATS ALIGNMENT MATCH SCORE", ln=0)
    
    pdf.set_xy(27, pdf.get_y() + 10)
    pdf.set_font("Helvetica", "I", 8.5)
    pdf.set_text_color(120, 130, 140)
    pdf.cell(100, 5, "Optimized for high-dimensional technical terms and requirements.", ln=0)
    
    pdf.set_xy(135, pdf.get_y() - 5)
    pdf.set_font("Helvetica", "B", 24)
    pdf.set_text_color(*accent_rgb)
    pdf.cell(50, 15, f"{res.match_score}%", ln=1, align="R")
    
    # Footer disclaimer on cover page
    pdf.set_xy(20, 265)
    pdf.set_font("Helvetica", "I", 7.5)
    pdf.set_text_color(156, 163, 175)
    pdf.cell(0, 5, "Confidential Document Package. Developed in accordance with advanced resume and application guidelines.", ln=1, align="C")
    
    # ----------------------------------------------------
    # PAGE 2: EXECUTIVE ALIGNMENT & STRATEGY
    # ----------------------------------------------------
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(26, 86, 219)
    pdf.cell(0, 10, "Executive Alignment & Strategic Summary", ln=1)
    pdf.set_draw_color(26, 86, 219)
    pdf.set_linewidth(0.8)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)
    
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(17, 24, 39)
    pdf.cell(0, 8, "EXECUTIVE FIT SUMMARY", ln=1)
    pdf.ln(1)
    
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(55, 65, 81)
    pdf.multi_cell(0, 5.5, clean_for_pdf(res.fit_summary))
    pdf.ln(10)
    
    # Direct Matches & Gap Areas
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(17, 24, 39)
    pdf.cell(0, 8, "ATS KEY ALIGNMENTS & KEY GAP UPDATES", ln=1)
    pdf.ln(1)
    
    pdf.set_font("Helvetica", "", 9.5)
    pdf.set_text_color(55, 65, 81)
    
    # Render gaps/suggestions using custom renderer
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_text_color(26, 86, 219)
    pdf.cell(0, 7, "Key Recommendations to Address Identified Gaps:", ln=1)
    pdf.ln(1)
    
    for sug in res.resume_suggestions:
        pdf.set_font("Helvetica", "", 9.5)
        pdf.set_text_color(31, 41, 55)
        pdf.cell(5, 5.5, "-", ln=0)
        
        old_l_margin = pdf.l_margin
        pdf.set_left_margin(old_l_margin + 5)
        write_rich_text(pdf, sug, default_font="Helvetica", default_size=9.5, default_style="", default_color=(31, 41, 55), line_height=5.5)
        pdf.ln(5.5)
        pdf.set_left_margin(old_l_margin)
        pdf.ln(1)
        
    # ----------------------------------------------------
    # PAGE 3: TAILORED COVER LETTER
    # ----------------------------------------------------
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(26, 86, 219)
    pdf.cell(0, 10, "1. Tailored Cover Letter", ln=1)
    pdf.set_draw_color(26, 86, 219)
    pdf.set_linewidth(0.8)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)
    
    write_markdown_line_by_line(pdf, res.cover_letter)
    
    # ----------------------------------------------------
    # PAGE 4: OPTIMIZED PROFESSIONAL RESUME
    # ----------------------------------------------------
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(26, 86, 219)
    pdf.cell(0, 10, "2. Optimized Professional Resume", ln=1)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)
    
    write_markdown_line_by_line(pdf, res.tailored_resume)
    
    # ----------------------------------------------------
    # PAGE 5: STRATEGIC INTERVIEW PREPARATION
    # ----------------------------------------------------
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(26, 86, 219)
    pdf.cell(0, 10, "3. Targeted Interview Preparation", ln=1)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(5)
    
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(17, 24, 39)
    pdf.cell(0, 8, "Likely Interview Questions & Core Talking Points:", ln=1)
    pdf.ln(2)
    
    for i, prep in enumerate(res.interview_prep):
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(26, 86, 219)
        pdf.cell(0, 6, f"Question {i+1}:", ln=1)
        pdf.ln(1)
        
        old_l_margin = pdf.l_margin
        pdf.set_left_margin(old_l_margin + 5)
        
        write_markdown_line_by_line(pdf, prep)
        pdf.ln(4)
        pdf.set_left_margin(old_l_margin)
        
    return bytes(pdf.output())

def extract_keywords_from_linkedin_url(url: str) -> list[str]:
    if not url:
        return []
    import re
    path_part = ""
    if "/in/" in url:
        path_part = url.split("/in/")[-1]
    elif "/pub/" in url:
        path_part = url.split("/pub/")[-1]
    else:
        path_part = url.split("/")[-1] or url.split("/")[-2] if len(url.split("/")) > 1 else url
    
    path_part = path_part.split("?")[0].split("#")[0]
    
    words = re.findall(r'[a-zA-Z0-9]+', path_part)
    filtered_words = []
    for w in words:
        w_lower = w.lower()
        if len(w_lower) > 2 and not w_lower.isdigit():
            if w_lower not in ['linkedin', 'com', 'www', 'http', 'https', 'in', 'pub', 'profile']:
                filtered_words.append(w_lower)
    return filtered_words

def extract_linkedin_profile_details(url: str) -> dict:
    if not url:
        return {
            "name": "Kevin Scott Pinard",
            "headline": "AI Developer, Data Analyst, & Automation Architect",
            "location": "Orlando, FL",
            "skills": ["Python", "AI", "SQL", "BigQuery", "Data Analysis", "Automation"],
            "raw_slug": ""
        }
    
    import re
    # Extract path portion
    path_part = ""
    if "/in/" in url:
        path_part = url.split("/in/")[-1]
    elif "/pub/" in url:
        path_part = url.split("/pub/")[-1]
    else:
        path_part = url.split("/")[-1] or url.split("/")[-2] if len(url.split("/")) > 1 else url
        
    path_part = path_part.split("?")[0].split("#")[0].strip()
    
    # Normalize path portion
    normalized = re.sub(r'[_/]', '-', path_part)
    parts = [p.strip() for p in normalized.split("-") if p.strip()]
    
    # Filter out numeric or hexadecimal IDs at the end
    cleaned_parts = []
    for p in parts:
        if re.match(r'^[0-9a-fA-F]+$', p) and (len(p) >= 6 or p.isdigit()):
            continue
        cleaned_parts.append(p)
        
    if not cleaned_parts:
        cleaned_parts = ["pinard", "kevin"] # Fallback if empty slug
        
    # Lookups for classification
    known_locations = {
        "orlando": "Orlando, FL",
        "nyc": "New York, NY",
        "newyork": "New York, NY",
        "sf": "San Francisco, CA",
        "sanfrancisco": "San Francisco, CA",
        "london": "London, UK",
        "austin": "Austin, TX",
        "seattle": "Seattle, WA",
        "boston": "Boston, MA",
        "chicago": "Chicago, IL",
        "miami": "Miami, FL",
        "atlanta": "Atlanta, GA",
        "florida": "Florida, USA",
        "california": "California, USA",
        "texas": "Texas, USA"
    }
    
    known_roles = {
        "developer": "Developer",
        "engineer": "Engineer",
        "scientist": "Scientist",
        "architect": "Architect",
        "analyst": "Analyst",
        "consultant": "Consultant",
        "manager": "Manager",
        "lead": "Lead",
        "specialist": "Specialist",
        "expert": "Expert"
    }
    
    known_skills = {
        "python": "Python",
        "ai": "AI",
        "ml": "Machine Learning",
        "data": "Data Science/Analytics",
        "cloud": "Cloud Computing",
        "aws": "AWS",
        "sql": "SQL",
        "react": "React",
        "java": "Java",
        "automation": "Automation",
        "agentic": "Agentic AI",
        "agent": "AI Agents",
        "bi": "Business Intelligence",
        "devops": "DevOps",
        "tableau": "Tableau",
        "bigquery": "BigQuery",
        "scrum": "Agile/Scrum",
        "pm": "Project Management",
        "fullstack": "Full Stack",
        "frontend": "Frontend",
        "backend": "Backend"
    }
    
    name_parts = []
    location_list = []
    headline_parts = []
    skills_list = []
    
    for i, word in enumerate(cleaned_parts):
        w_lower = word.lower()
        if w_lower in known_locations:
            location_list.append(known_locations[w_lower])
        elif w_lower in known_roles:
            headline_parts.append(known_roles[w_lower])
            skills_list.append(known_roles[w_lower])
        elif w_lower in known_skills:
            skills_list.append(known_skills[w_lower])
            headline_parts.append(known_skills[w_lower])
        else:
            if len(name_parts) < 3 and not any(char.isdigit() for char in word):
                name_parts.append(word.title())
            else:
                skills_list.append(word.title())
                
    if len(name_parts) == 1:
        name_str = name_parts[0]
        if name_str.lower() == "pinardkevin":
            name_parts = ["Kevin", "Scott", "Pinard"]
        elif name_str.lower() == "kevinpinard":
            name_parts = ["Kevin", "Pinard"]
        else:
            name_parts = [name_str.title()]
            
    name = " ".join(name_parts) if name_parts else "Kevin Scott Pinard"
    location = location_list[0] if location_list else "Orlando, FL"
    
    if headline_parts:
        seen = set()
        clean_headline_parts = []
        for hp in headline_parts:
            if hp not in seen:
                seen.add(hp)
                clean_headline_parts.append(hp)
        headline = " | ".join(clean_headline_parts)
    else:
        headline = "AI Developer, Data Analyst, & Automation Architect"
        
    if not skills_list:
        skills_list = ["Python", "AI", "SQL", "BigQuery", "Data Analysis", "Automation"]
        
    skills_list = list(dict.fromkeys(skills_list))
    
    return {
        "name": name,
        "headline": headline,
        "location": location,
        "skills": skills_list,
        "raw_slug": path_part
    }

def get_fresh_gcloud_token() -> str:
    import subprocess
    try:
        token = subprocess.check_output(
            ["gcloud", "auth", "print-access-token"],
            text=True,
            shell=True,
            stderr=subprocess.DEVNULL
        ).strip()
        if token:
            return token
    except Exception:
        pass
    return None

# Load local .env
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="EasyApplier ◆ AI Agent Dashboard",
    page_icon="◆",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Theme selection (default to light for cornsilk theme)
if "theme" not in st.session_state:
    st.session_state.theme = "light"

def toggle_theme():
    st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"

IS_DARK = st.session_state.theme == "dark"

# CSS variables depending on theme
if IS_DARK:
    bg_color = "#09090b"
    bg_subtle = "#0c0c0f"
    card_color = "#0c0c0f"
    card_hover = "#131316"
    border_color = "#1e1e24"
    border_subtle = "#16161a"
    text_color = "#fafafa"
    text_muted = "#71717a"
    text_dim = "#52525b"
    shadow = "none"
    accent_color = "#6366f1"
    accent_hover = "#4f46e5"
else:
    bg_color = "#ffffff" # Pure White background
    bg_subtle = "#f6f8fa" # Very light grey sidebar/subtle background
    card_color = "#ffffff" 
    card_hover = "#f6f8fa" 
    border_color = "#d0d7de" # Soft grey borders
    border_subtle = "#ebf0f4"
    text_color = "#111111" # High-contrast near-black text
    text_muted = "#4b5563" # Soft charcoal grey for muted notes/labels
    text_dim = "#6b7280" # Lighter charcoal
    shadow = "0 1px 3px rgba(0,0,0,0.05), 0 1px 2px rgba(0,0,0,0.05)"
    accent_color = "#4285F4" # Google Logo Blue
    accent_hover = "#357ae8" # Darker Google Blue

# Inject custom CSS matching shared design patterns
st.markdown(f"""
<style>
    :root {{
        --bg: {bg_color};
        --bg-subtle: {bg_subtle};
        --card: {card_color};
        --card-hover: {card_hover};
        --border: {border_color};
        --border-subtle: {border_subtle};
        --text: {text_color};
        --text-muted: {text_muted};
        --text-dim: {text_dim};
        --accent: {accent_color};
        --accent-hover: {accent_hover};
        --green: #10b981;
        --green-muted: rgba(16, 185, 129, 0.1);
        --red: #ef4444;
        --red-muted: rgba(239, 68, 68, 0.1);
        --amber: #f59e0b;
        --amber-muted: rgba(245, 158, 11, 0.1);
        --shadow: {shadow};
        --radius: 10px;
    }}

    /* Hide default elements */
    header[data-testid="stHeader"], #MainMenu, footer, [data-testid="stToolbar"],
    [data-testid="stDecoration"], [data-testid="stStatusWidget"], .stDeployButton {{
        display: none !important;
    }}

    /* Global styling overrides */
    html, body, [data-testid="stAppViewContainer"], [data-testid="stApp"], .main, .block-container, section[data-testid="stMain"], section[data-testid="stSidebar"] {{
        background-color: var(--bg) !important;
        color: var(--text) !important;
        font-family: 'Plus Jakarta Sans', -apple-system, sans-serif !important;
    }}
    section[data-testid="stSidebar"] {{
        background-color: var(--bg-subtle) !important;
        border-right: 1px solid var(--border) !important;
    }}
    
    /* Style inputs, textareas, and select boxes specifically to prevent lighter overrides */
    .stTextInput input, .stTextArea textarea, 
    div[data-baseweb="base-input"] input, div[data-baseweb="textarea"] textarea,
    div[data-baseweb="select"] div, span[data-baseweb="select"],
    div[role="listbox"] div, div[role="option"] p, div[role="option"] span,
    div[data-testid="stRadio"] label, div[data-testid="stRadio"] label span, div[data-testid="stRadio"] label p,
    div[data-testid="stFileUploader"] section, div[data-testid="stFileUploader"] p, div[data-testid="stFileUploader"] span,
    .stSelectbox div[role="button"] {{
        color: var(--text) !important;
        -webkit-text-fill-color: var(--text) !important;
        background-color: var(--card) !important;
    }}

    /* Inputs/Textareas container style */
    div[data-baseweb="base-input"], div[data-baseweb="textarea"] {{
        border: 1px solid var(--border) !important;
        background-color: var(--card) !important;
    }}

    /* Placeholders in inputs */
    .stTextInput input::placeholder, .stTextArea textarea::placeholder {{
        color: var(--text-dim) !important;
        opacity: 0.75 !important;
        -webkit-text-fill-color: var(--text-dim) !important;
    }}

    /* Text elements & widgets labels */
    div[data-testid="stWidgetLabel"] p, label[data-testid="stWidgetLabel"] p,
    .stMarkdown p, .stMarkdown li, .stMarkdown ul, .stMarkdown ol, 
    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4, .stMarkdown h5, .stMarkdown h6,
    div[data-testid="stRadio"] p, label[data-baseweb="radio"] div,
    [data-testid="stMarkdownContainer"] p, [data-testid="stMarkdownContainer"] li, [data-testid="stMarkdownContainer"] span {{
        color: var(--text) !important;
    }}

    /* Keep button labels clean (override general span modifications) */
    div[data-testid="stButton"] button span, .stButton button span {{
        color: #ffffff !important;
    }}

    .block-container {{
        padding: 2rem 2.5rem 3rem !important;
        max-width: 1360px !important;
    }}

    /* Custom Metric Cards */
    .metric-card {{
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1.25rem 1.4rem;
        box-shadow: var(--shadow);
    }}
    .metric-label {{
        font-size: 0.8rem;
        color: var(--text-muted);
        font-weight: 500;
        margin-bottom: 0.2rem;
    }}
    .metric-value {{
        font-size: 1.6rem;
        font-weight: 700;
        color: var(--text);
        letter-spacing: -0.02em;
    }}

    /* Custom Chart wrap */
    .chart-wrap {{
        background: var(--card);
        border: 1px solid var(--border);
        border-radius: var(--radius);
        padding: 1.2rem;
        box-shadow: var(--shadow);
        margin-bottom: 1.5rem;
    }}
    .chart-title {{
        font-size: 0.9rem;
        font-weight: 600;
        color: var(--text);
    }}
    .chart-subtitle {{
        font-size: 0.75rem;
        color: var(--text-dim);
        margin-bottom: 1rem;
    }}

    /* Horizontal list gap */
    [data-testid="stHorizontalBlock"] {{
        gap: 1.25rem !important;
    }}

    /* Custom pill-style tabs */
    button[data-baseweb="tab"] {{
        background: transparent !important;
        color: var(--text-muted) !important;
        font-size: 0.85rem !important;
        font-weight: 500 !important;
        padding: 0.6rem 1.2rem !important;
        border: 1px solid transparent !important;
        border-radius: 7px !important;
        color: var(--text-dim) !important;
    }}
    button[data-baseweb="tab"][aria-selected="true"] {{
        color: var(--text) !important;
        background: var(--card) !important;
        border-color: var(--border) !important;
    }}
    [data-baseweb="tab-highlight"], [data-baseweb="tab-border"] {{
        display: none !important;
    }}
    [data-baseweb="tab-list"] {{
        gap: 6px !important;
        background: var(--bg-subtle) !important;
        border: 1px solid var(--border) !important;
        border-radius: 10px !important;
        padding: 4px;
        margin-bottom: 20px;
    }}

    /* Badges */
    .badge {{
        display: inline-block;
        padding: 3px 10px;
        border-radius: 6px;
        font-size: 0.75rem;
        font-weight: 500;
    }}
    .badge-green {{ color: var(--green); background: var(--green-muted); border: 1px solid rgba(16, 185, 129, 0.2); }}
    .badge-red {{ color: var(--red); background: var(--red-muted); border: 1px solid rgba(239, 68, 68, 0.2); }}
    .badge-amber {{ color: var(--amber); background: var(--amber-muted); border: 1px solid rgba(245, 158, 11, 0.2); }}
    .badge-blue {{ color: var(--accent); background: rgba(99, 102, 241, 0.1); border: 1px solid rgba(99, 102, 241, 0.2); }}

    /* Style all buttons on the page to use brand blue */
    div[data-testid="stButton"] button, .stButton button {{
        background-color: var(--accent) !important;
        color: #ffffff !important;
        border: 1px solid var(--accent) !important;
        transition: background-color 0.2s, border-color 0.2s !important;
    }}
    div[data-testid="stButton"] button:hover, .stButton button:hover {{
        background-color: var(--accent-hover) !important;
        border-color: var(--accent-hover) !important;
        color: #ffffff !important;
    }}
</style>
""", unsafe_allow_html=True)

# Custom helpers
def extract_text_from_file(uploaded_file):
    if uploaded_file is None:
        return ""
    try:
        file_bytes = uploaded_file.read()
        filename = uploaded_file.name.lower()
        if filename.endswith(".docx"):
            doc = docx.Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        elif filename.endswith(".pdf"):
            reader = pypdf.PdfReader(BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
            return text
        elif filename.endswith(".doc"):
            # Basic fallback for legacy doc files
            decoded = file_bytes.decode('utf-8', errors='ignore')
            lines = [line.strip() for line in decoded.split('\n') if len(line.strip()) > 3]
            return "\n".join(lines)
    except Exception as e:
        return f"Error parsing file: {e}"
    return ""

def load_candidate_profile() -> str:
    profile_parts = []
    path1 = "C:/Users/pinar/Kaggle-AI-Agents-Course/Google Kaggle 5DAYAI Vibe Coding Project/RESUME.Google.Day3.docx"
    path2 = "C:/Users/pinar/Google Final Submission Materials/Resume Base Build.docx"
    
    if os.path.exists(path1):
        try:
            doc = docx.Document(path1)
            t = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if t:
                profile_parts.append(f"--- CANDIDATE DAY 3 PROFILE ---\n{t}")
        except Exception:
            pass
            
    if os.path.exists(path2):
        try:
            doc = docx.Document(path2)
            t = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if t:
                profile_parts.append(f"--- CANDIDATE BASE BUILD PROFILE ---\n{t}")
        except Exception:
            pass
            
    if profile_parts:
        return "\n\n".join(profile_parts)
    return "Kevin Scott Pinard | Kevinpolymath@gmail.com | 352-406-3847 | Orlando, FL"

def render_metric_card(label, value):
    st.markdown(f"""
    <div class="metric-card">
        <div class="metric-label">{label}</div>
        <div class="metric-value">{value}</div>
    </div>
    """, unsafe_allow_html=True)

# Plotly theme configuration
PLOT_LAYOUT = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(family="Plus Jakarta Sans, sans-serif", color="#000000" if not IS_DARK else "#a1a1aa", size=11),
    margin=dict(l=40, r=20, t=20, b=40),
    xaxis=dict(
        gridcolor="rgba(0,0,0,0.05)" if not IS_DARK else "rgba(255,255,255,0.05)",
        zerolinecolor="rgba(0,0,0,0.05)" if not IS_DARK else "rgba(255,255,255,0.05)",
        tickfont=dict(size=10, color="#000000" if not IS_DARK else "#71717a"),
    ),
    yaxis=dict(
        gridcolor="rgba(0,0,0,0.05)" if not IS_DARK else "rgba(255,255,255,0.05)",
        zerolinecolor="rgba(0,0,0,0.05)" if not IS_DARK else "rgba(255,255,255,0.05)",
        tickfont=dict(size=10, color="#000000" if not IS_DARK else "#71717a"),
    ),
)

# Header Section
logo_path = "aligned_labs_logo.jpg"
if not os.path.exists(logo_path):
    logo_path = "C:/Users/pinar/.gemini/antigravity-cli/brain/5ab6de28-73cb-4611-8a6d-8d4c0c72d395/aligned_labs_logo_1781851225345.jpg"

if os.path.exists(logo_path):
    st.image(logo_path, width=150)

st.markdown("""
<div style="display:flex; align-items:center; gap:12px;">
    <span style="font-family:'Outfit', sans-serif; font-size:28px; font-weight:700; color:var(--accent);">EasyApplier</span>
    <span class="badge badge-green" style="font-size:12px; font-weight:600; border-radius:20px;">AI Dashboard</span>
</div>
""", unsafe_allow_html=True)
if "searched_title" in st.session_state and st.session_state.job_listings:
    st.markdown(f"<div style='font-size:14px; color:var(--text-muted); margin-top:4px;'>Most Recent {st.session_state.searched_title} 20 Job Postings on LinkedIn</div>", unsafe_allow_html=True)

st.markdown("<hr style='margin:1.2rem 0; border:0; border-top:1px solid var(--border);'>", unsafe_allow_html=True)

# Sidebar Configuration / Credentials
with st.sidebar:
    st.markdown("### ⚙️ Agent Settings")
    
    # Preloaded key defaults
    default_key = os.getenv("GEMINI_API_KEY", "")
    # If the user input is present in session state, override
    if "api_key" not in st.session_state:
        st.session_state.api_key = default_key if default_key else "AQ.Ab8RN6KUiT8oh9KCbfdN25iizco0cZkWN49mjwwetIaR4yRDiA"
        
    api_key_input = st.text_input(
        "Gemini API Key / Project Token",
        value=st.session_state.api_key,
        type="password",
        help="Paste your Google AI Studio API Key or course token starting with AIzaSy... or AQ.Ab..."
    )
    st.session_state.api_key = api_key_input

    model_input = st.selectbox(
        "Gemini Model",
        options=["gemini-2.5-flash", "gemini-2.5-pro", "gemini-2.0-flash-exp"],
        index=0
    )

# Load resume dynamically from local file to prevent manual upload/exposure of PII data
resume_input = load_candidate_profile()

# Pydantic schema for structured output
class ApplicationStrategy(BaseModel):
    match_score: int = Field(..., description="Overall fit score from 1 to 100")
    fit_summary: str = Field(..., description="A 3-4 sentence summary of why the candidate fits this role")
    cover_letter: str = Field(..., description="A professional, tailored cover letter based on the job and resume")
    tailored_resume: str = Field(..., description="A rewritten version of the candidate resume, fully optimized to align with the job description keywords")
    resume_suggestions: list[str] = Field(..., description="Bullet points suggesting specific updates to the resume to align with keywords")
    interview_prep: list[str] = Field(..., description="Top 3 likely interview questions with suggested talking points")

# Initialize strategy cache in session state
if "cached_strategies" not in st.session_state:
    st.session_state.cached_strategies = {}

# Main Content Layout
with st.form("search_form", border=False):
    search_col1, search_col2 = st.columns([6, 2])
    with search_col1:
        job_title = st.text_input("Search Job Title", placeholder="e.g. Python Developer, Machine Learning Engineer...", label_visibility="collapsed")
    with search_col2:
        search_clicked = st.form_submit_button("Search LinkedIn Jobs", type="primary", use_container_width=True)

# Retrieve job listings or use session state
if "job_listings" not in st.session_state:
    st.session_state.job_listings = []

if search_clicked and job_title:
    with st.spinner("Scraping 20 most recent job listings from LinkedIn guest search via API backend..."):
        try:
            # Call backend endpoint /api/jobs to fetch recent LinkedIn jobs
            backend_url = os.getenv("BACKEND_URL", "http://127.0.0.1:8000")
            response = requests.get(f"{backend_url}/api/jobs", params={"title": job_title}, timeout=25)
            if response.status_code == 200:
                st.session_state.job_listings = response.json()
                st.session_state.searched_title = job_title.title()
                st.toast(f"Successfully loaded {len(st.session_state.job_listings)} jobs!", icon="✅")
            else:
                try:
                    detail = response.json().get("detail", response.text)
                except Exception:
                    detail = response.text
                st.error(f"Error scraping LinkedIn via API: {detail}")
        except Exception as e:
            st.error(f"Failed to connect to backend job search API: {e}")

# If we have listings, render dashboard
if st.session_state.job_listings:
    df = pd.DataFrame(st.session_state.job_listings)
    
    # Calculate some helper metrics
    total_jobs = len(df)
    unique_companies = df['company'].nunique()
    top_location = df['location'].value_counts().index[0] if not df.empty else "N/A"
    
    # Metric KPI Row
    m1, m2, m3 = st.columns(3)
    with m1:
        render_metric_card("Total Jobs Loaded", f"{total_jobs} Postings")
    with m2:
        render_metric_card("Unique Companies Hiring", f"{unique_companies} Companies")
    with m3:
        render_metric_card("Primary Hiring Location", top_location)
        
    st.markdown("<br>", unsafe_allow_html=True)
    # Split main layout into 3 columns: Listings & Selection, Candidate Profile & Actions, AI Application Strategy
    col_list, col_profile, col_tailor = st.columns([4.5, 3, 4.5])
    
    with col_list:
        st.markdown("### 💼 Job Listings")
        
        # 3. Add filters for "Easy Apply" and Workplace type
        f_col1, f_col2 = st.columns([1, 1.2])
        with f_col1:
            easy_apply_only = st.checkbox("🟢 Easy Apply Only", value=False)
        with f_col2:
            workplace_choices = st.multiselect(
                "Workplace Type",
                options=["Remote", "Hybrid", "Onsite"],
                default=["Remote", "Hybrid", "Onsite"],
                label_visibility="collapsed"
            )
            
        # Filter dataframe
        filtered_df = df.copy()
        if easy_apply_only:
            if 'easy_apply' in filtered_df.columns:
                filtered_df = filtered_df[filtered_df['easy_apply'] == True]
        if workplace_choices:
            if 'workplace_type' in filtered_df.columns:
                filtered_df = filtered_df[filtered_df['workplace_type'].isin(workplace_choices)]
                
        if filtered_df.empty:
            st.warning("No job listings match your selected filters.")
            selected_job = None
        else:
            job_options = []
            job_map = {}
            for idx, row in filtered_df.iterrows():
                # Add Easy Apply badge if applicable
                easy_badge = " [Easy Apply]" if row.get('easy_apply') else ""
                work_badge = f" [{row.get('workplace_type')}]" if row.get('workplace_type') else ""
                opt_key = f"{row['title']} @ {row['company']}{easy_badge}{work_badge}"
                job_options.append(opt_key)
                job_map[opt_key] = idx
                
            selected_option = st.radio("Select job posting to optimize:", options=job_options, label_visibility="collapsed")
            selected_idx = job_map[selected_option]
            
            # State-change detection: Clear stale results if the user selects a new job listing
            if "previous_selected_idx" not in st.session_state:
                st.session_state.previous_selected_idx = selected_idx
            elif st.session_state.previous_selected_idx != selected_idx:
                st.session_state.previous_selected_idx = selected_idx
                if "tailor_result" in st.session_state:
                    del st.session_state.tailor_result
                    
            selected_job = st.session_state.job_listings[selected_idx]
            st.session_state.selected_job = selected_job
            
            # Scrape and cache job description dynamically
            if "scraped_desc_cache" not in st.session_state:
                st.session_state.scraped_desc_cache = {}
                
            job_link = selected_job['link']
            if job_link not in st.session_state.scraped_desc_cache:
                with st.spinner("Fetching full job description from LinkedIn..."):
                    try:
                        desc = get_linkedin_job_description(job_link)
                        st.session_state.scraped_desc_cache[job_link] = desc
                    except Exception as ex:
                        st.session_state.scraped_desc_cache[job_link] = f"Error fetching details: {ex}"
                        
            full_desc = st.session_state.scraped_desc_cache.get(job_link, "")
            
            # Display detailed card
            st.markdown(f"""
            <div style="background:var(--card-hover); border:1px solid var(--border); border-radius:var(--radius); padding:1rem; margin-top:1rem;">
                <h4 style="margin:0 0 4px 0;">{selected_job['title']}</h4>
                <div style="color:var(--text-muted); font-size:13px; margin-bottom:8px;">{selected_job['company']} — {selected_job['location']}</div>
                <div class="badge badge-blue">{selected_job['posted_date']}</div>
                {"<div class='badge badge-green' style='margin-left:5px; display:inline-block;'>Easy Apply</div>" if selected_job.get('easy_apply') else ""}
                <div class="badge badge-amber" style="margin-left:5px; display:inline-block;">{selected_job.get('workplace_type', 'Onsite')}</div>
                <p style="margin-top:10px;"><a href="{selected_job['link']}" target="_blank" style="color:var(--accent); text-decoration:none; font-weight:600;">View original post on LinkedIn ↗</a></p>
            </div>
            """, unsafe_allow_html=True)
            
            if full_desc and not full_desc.startswith("Error fetching"):
                st.markdown("<br>", unsafe_allow_html=True)
                with st.expander("📝 Scraped Job Description", expanded=False):
                    st.text_area("Full Description", value=full_desc, height=250, disabled=True, label_visibility="collapsed")

    with col_profile:
        st.markdown("### 👤 Candidate Details")
        
        linkedin_url = st.text_input("1. Enter LinkedIn Profile URL", value="https://www.linkedin.com/in/pinardkevin", placeholder="https://www.linkedin.com/in/username")
        
        # Extract rich candidate details
        profile_details = extract_linkedin_profile_details(linkedin_url)
        
        st.markdown(f"""
        <div style="background:var(--card-hover); border:1px solid var(--border); border-radius:var(--radius); padding:1.2rem; margin-bottom:1rem;">
            <div style="display:flex; justify-content:space-between; align-items:center;">
                <h4 style="margin:0; font-size:15px;">✨ Extracted Profile Details</h4>
                <span class="badge badge-blue" style="font-size:10px;">Slug: {profile_details['raw_slug']}</span>
            </div>
            <div style="margin-top:10px; font-size:13px; line-height:1.5;">
                <p style="margin:4px 0;"><strong>Name:</strong> {profile_details['name']}</p>
                <p style="margin:4px 0;"><strong>Headline:</strong> {profile_details['headline']}</p>
                <p style="margin:4px 0;"><strong>Location:</strong> {profile_details['location']}</p>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("**Inferred Skills & Roles:**")
        pills_html = "".join([f"<span class='badge badge-blue' style='margin-right: 5px; margin-bottom: 5px; display: inline-block;'>{kw}</span>" for kw in profile_details['skills']])
        st.markdown(f"<div style='margin-bottom: 15px;'>{pills_html}</div>", unsafe_allow_html=True)
        
        # Load local base profile
        default_resume = load_candidate_profile()
        
        # Build richer candidate profile template combining extracted profile and fallback resume
        resume_text_doc = f"""
        =====================================================
        EXTRACTED CANDIDATE PROFILE (LINKEDIN URL)
        =====================================================
        Full Name: {profile_details['name']}
        Professional Headline: {profile_details['headline']}
        Target Location: {profile_details['location']}
        Inferred Skills & Focus: {", ".join(profile_details['skills'])}
        
        =====================================================
        CANDIDATE BASE RESUME (FALLBACK & EXPERIENCE SOURCE)
        =====================================================
        {default_resume}
        
        CRITICAL PARSING GUIDELINES:
        - The candidate's name MUST be parsed as "{profile_details['name']}".
        - The candidate's summary and location MUST align with "{profile_details['location']}" and headline "{profile_details['headline']}".
        - Merge the inferred skills ({", ".join(profile_details['skills'])}) into the resume.
        - Ensure all tailored outputs use "{profile_details['name']}" as the candidate's name.
        """
                
        # Validate LinkedIn URL
        is_valid_linkedin = False
        if linkedin_url:
            cleaned_url = linkedin_url.strip().lower()
            if (cleaned_url.startswith("http://") or cleaned_url.startswith("https://")) and "linkedin.com" in cleaned_url:
                is_valid_linkedin = True
                
        cache_key = (
            selected_job['link'] if (selected_job and 'link' in selected_job) else "",
            linkedin_url,
            st.session_state.api_key
        )
        
        # Load from cache if available
        if cache_key in st.session_state.cached_strategies:
            st.session_state.tailor_result = st.session_state.cached_strategies[cache_key]
        else:
            if "tailor_result" in st.session_state:
                del st.session_state.tailor_result
                
        is_generated = cache_key in st.session_state.cached_strategies
        
        # Action button
        st.markdown("<br>", unsafe_allow_html=True)
        optimize_btn = st.button(
            "Generate Tailored Career Path",
            width="stretch",
            type="primary",
            disabled=is_generated or not (is_valid_linkedin and selected_job)
        )
        
        if not selected_job:
            st.caption("⚠️ Please select a Job Listing.")
        elif not linkedin_url:
            st.caption("⚠️ Please enter your LinkedIn Profile URL.")
        elif not is_valid_linkedin:
            st.caption("⚠️ Please enter a valid LinkedIn Profile URL (must start with http/https and contain 'linkedin.com').")
            
        if optimize_btn:
            if not st.session_state.api_key:
                st.warning("Please configure your Gemini API key in the sidebar.")
            else:
                with st.spinner("Initiating Multi-Agent Workflow... Researching job posting, aligning candidate profile, and auditing tailored drafts..."):
                    # Ensure we have the full job description
                    full_job_desc = st.session_state.scraped_desc_cache.get(selected_job['link'], "")
                    if not full_job_desc or full_job_desc.startswith("Error fetching"):
                        full_job_desc = f"Job Title: {selected_job['title']}\nCompany: {selected_job['company']}\nLocation: {selected_job['location']}"

                    try:
                        # Set active model in environment
                        os.environ["GEMINI_MODEL"] = model_input
                        
                        # Configure dynamic gcloud oauth credentials if starts with ya29 or AQ
                        key_to_use = st.session_state.api_key
                        if key_to_use and (key_to_use.startswith("ya29.") or key_to_use.startswith("AQ.")):
                            fresh_token = get_fresh_gcloud_token()
                            if fresh_token:
                                key_to_use = fresh_token
                                st.session_state.api_key = fresh_token
                        
                        if key_to_use:
                            os.environ["GEMINI_API_KEY"] = key_to_use

                        # Initialize and run the multi-agent orchestrator
                        orchestrator = EasyApplierOrchestrator()
                        result = orchestrator.run_workflow(
                            job_title=selected_job['title'],
                            raw_job_description=full_job_desc,
                            raw_resume_text=resume_text_doc
                        )
                        
                        # Map OrchestratorResult back to ApplicationStrategy for UI rendering compatibility
                        mapped_strategy = ApplicationStrategy(
                            match_score=result.gap_analysis.match_score,
                            fit_summary=result.gap_analysis.fit_summary,
                            cover_letter=result.tailor_output.cover_letter,
                            tailored_resume=result.tailor_output.tailored_resume,
                            resume_suggestions=result.gap_analysis.gap_areas,
                            interview_prep=result.tailor_output.interview_prep
                        )

                        st.session_state.cached_strategies[cache_key] = mapped_strategy
                        st.session_state.tailor_result = mapped_strategy
                        st.toast("Strategy generated successfully by multi-agent system!", icon="✨")
                        st.rerun()

                    except Exception as e:
                        err_msg = str(e).upper()
                        if "429" in err_msg or "RESOURCE_EXHAUSTED" in err_msg or "RATE_LIMIT" in err_msg:
                            st.error("Rate limit reached. Please wait a moment and try again.")
                        else:
                            st.error(f"Failed to generate strategy via multi-agent system: {e}")
                            
    with col_tailor:
        st.markdown("### 🏆 AI Application Strategy")
        
        # Render results if generated
        if "tailor_result" in st.session_state:
            res = st.session_state.tailor_result
            score = res.match_score
            badge_class = "badge-green" if score >= 80 else ("badge-amber" if score >= 60 else "badge-red")
            
            st.markdown(f"""
            <div style="background:rgba(99, 102, 241, 0.05); border:1px solid rgba(99, 102, 241, 0.2); border-radius:var(--radius); padding:1rem; margin-bottom:1.5rem; display:flex; justify-content:space-between; align-items:center;">
                <div>
                    <h4 style="margin:0 0 2px 0;">Resume Match Score</h4>
                    <div style="color:var(--text-muted); font-size:12px;">Based on keywords, skills, and experience alignment</div>
                </div>
                <div class="badge {badge_class}" style="font-size:1.6rem; font-weight:800; padding:8px 16px;">{score}%</div>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("#### 📝 Executive Fit Recommendation")
            st.write(res.fit_summary)
            st.markdown("<br>", unsafe_allow_html=True)
            
            # Generate FPDF document package
            with st.spinner("Generating beautiful tailored career path PDF document..."):
                try:
                    pdf_bytes = generate_strategy_pdf(res, selected_job['title'], selected_job['company'], linkedin_url)
                    
                    st.download_button(
                        label="📥 Download Tailored Career Package (.pdf)",
                        data=pdf_bytes,
                        file_name=f"EasyApplier_Tailored_Career_Package_{selected_job['company'].replace(' ', '_')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    st.success("Your highly optimized resume, cover letter, suggestions, and interview prep is compiled into the custom PDF above!")
                except Exception as pdf_ex:
                    st.error(f"Error compiling strategy PDF package: {pdf_ex}")
                    # Standard display fallback
                    with st.expander("📄 View Tailored Resume Draft"):
                        st.code(res.tailored_resume)
        else:
            st.info("Select a job listing on the left, verify your profile details, and click 'Generate Tailored Career Path' to view the optimized strategy.")

    # Raw scraped data view in collapsible bottom expander
    st.markdown("<br><hr style='margin:1.5rem 0; border:0; border-top:1px solid var(--border);'>", unsafe_allow_html=True)
    with st.expander("🔍 Raw Scraped Job Data View", expanded=False):
        st.dataframe(df, width="stretch")
        
else:
    st.markdown("""
    <div style="text-align:center; padding:80px 40px; border:1px dashed var(--border); border-radius:var(--radius); color:var(--text-muted);">
        <svg width="64" height="64" fill="none" stroke="currentColor" stroke-width="1.2" viewBox="0 0 24 24" style="margin-bottom:15px; color:var(--text-dim);">
            <path stroke-linecap="round" stroke-linejoin="round" d="M20.25 14.15v4.25c0 .621-.504 1.125-1.125 1.125H4.875c-.621 0-1.125-.504-1.125-1.125v-4.25m16.5 0a2.18 2.18 0 00.75-1.661V8.706c0-1.081-.768-2.015-1.837-2.175a48.114 48.114 0 00-3.413-.387m4.5 8.006c-.194.165-.45.258-.717.258H3.75c-.266 0-.523-.093-.717-.258m16.5 0a2.18 2.18 0 01.75-1.661V8.706c0-1.081-.768-2.015-1.837-2.175a48.114 48.114 0 00-3.413-.387m0 0V6.25c0-1.098-.826-1.97-1.922-2.025a47.95 47.95 0 00-6.071 0C5.122 4.28 4.3 5.152 4.3 6.25v.806m9.75 0a48.693 48.693 0 01-6.071 0"></path>
        </svg>
        <h3>No Job Listings Loaded</h3>
        <p style="font-size:14px; margin-top:5px;">Enter a job title in the search box above to retrieve live postings from LinkedIn guest search.</p>
    </div>
    """, unsafe_allow_html=True)


